from docx import Document
from docx.enum.style import WD_STYLE_TYPE  # style type
from docx.shared import Pt, RGBColor, Cm, Inches  # format para
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # para alignment
from colorama import Fore, Style  # adding colour to command line
from docx.oxml.shared import OxmlElement # for adding page background
from docx.oxml.ns import qn  # for adding page background
from docx.enum.section import WD_SECTION_START  # adding section
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # Table Format
import matplotlib.pyplot as plt  # for graphs
import openai
import numpy as np #for graphs
import pandas as pd
import os

doc = Document()
print(Fore.BLUE + Style.BRIGHT +'''Hey want to create a Word document of you annual sales report
Dont worry we got you!!!
Just type the required fields below.''')
print()

openai.api_key = "" #Add your own OpenAI key

# adding background to the document
def add_background():
    # Now Add below children to root xml tree
    # creating xml element using Oxml Element
    shd = OxmlElement('w:background')
    # Adding attributes to the xml element
    shd.set(qn('w:color'), '0D0D0D')  # black color
    shd.set(qn('w:themeColor'), 'text1')
    shd.set(qn('w:themeTint'), 'F2')
    # Adding background element at the start of Document.xml using below
    doc.element.insert(0, shd)
    # Adding displayBackgroundShape element to setting.xml
    shd1 = OxmlElement('w:displayBackgroundShape')
    doc.settings.element.insert(0, shd1)


#creating sections in the document
def create_sections():
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


                                           # Heading 1 - COMPANY HISTORY

# creating section 1
def create_H1_Para1():
    def error_management1():
        while True:
            try:
                H1 = input(Fore.YELLOW + Style.BRIGHT + 'Enter the name of the company - ')
                if H1.isalpha() or ' ' in H1:
                    pass
                else:
                    raise ValueError
                return H1
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    H1 = error_management1()

    # adding styles to the heading
    def H1_style():
        head1 = doc.add_heading(H1, 7)
        font = head1.style.font
        font.name = 'Cambria'
        font.size = Pt(48)
        font.bold = True
        font.italic = False
        font.color.rgb = RGBColor(160, 160, 160)
        head1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    H1_style()

    #  Generating a response using chatGPT
    def chatgpt1():
        prompt = f"give the history of the company {H1}. If the mentioned company is fake create a factual history. if the company exists in the real word give a real history. divide the output into 4 paragraphs."
        completion = openai.Completion.create(engine='text-davinci-003', prompt=prompt, max_tokens=1024,  n=3, stop=None,temperature=0)  # controls the level of randomness in the output. The higher the temperature, the more varied and less coherent response output
        response1 = completion.choices[0].text
        return response1
    response1 = chatgpt1()

    # setting the style for para1
    def p1_style():
        p1 = doc.add_paragraph(response1)
        font1 = p1.style.font
        font1.name = 'Cascadia Code'
        font1.size = Pt(12)
        font1.color.rgb = RGBColor(255, 255, 255)
        p1_format = p1.paragraph_format
        p1_format.space_before = Pt(0)
        p1_format.space_after = Pt(0)
        p1_format.keep_together = True
        p1_format.keep_with_next = False
        p1_format.page_break_before = False
        p1_format.widow_control = False
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_page_break()
    p1_style()
    return H1
H1 = create_H1_Para1()
# Heading1 function complete


                                         # Heading 2 -  COMPANY PERFORMANCE

# creating section 2
def create_H2_Para2():

    # adding styles to H2
    def H2_style():
        H2 = 'Performance'
        head2 = doc.add_heading(H2, 6)
        font2 = head2.style.font
        font2.size = Pt(48)
        font2.bold = True
        font2.italic = False
        font2.color.rgb = RGBColor(160, 160, 160)
        head2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(2):
            doc.add_paragraph(' ')
    H2_style()

                                               # Para 2[Table and Graph]
    print()
    print(Fore.RED + Style.BRIGHT + '#Always write the input in the form(___,____etc) in the order unless mentioned')
    print()
    print(Fore.YELLOW + Style.BRIGHT + 'Table - 1')

    # taking input of how many rows and columns does the user want
    def error_management2():
        while True:
            try:
                row_col1 = input(Fore.MAGENTA + Style.BRIGHT + 'How many rows and columns do you want - ')  # taking the order of the table
                new_row_col1 = row_col1.split(',')
                if len(new_row_col1) == 2:
                    for i in new_row_col1:
                        if i.isdigit():
                            pass
                        else:
                            raise ValueError
                else:
                    raise TypeError
                return new_row_col1
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    new_row_col1 = error_management2()
    rows1 = int(new_row_col1[0])
    cols1 = int(new_row_col1[1])
    print()

    # taking input of the column titles
    def error_management3():
        while True:
            try:
                value1 = input(Fore.MAGENTA + Style.BRIGHT + 'Input the column Titles - ')
                valuesplit1 = value1.split(',')
                if len(valuesplit1) == cols1:
                    for i in valuesplit1:
                        if i.replace(' ', '').replace('.', '').replace('%', '').replace('(', '').replace(')', '').replace('$', '').isalpha():
                            pass
                        else:
                            raise ValueError
                else:
                    raise TypeError
                return valuesplit1
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    valuesplit1 = error_management3()
    print()

    # taking input of the column values
    def error_management4():
        records1 = []
        for i in range(1, rows1 + 1):  # taking table values from the user
            def tried():
                while True:
                    try:
                        t = input(Fore.MAGENTA + Style.BRIGHT + f'Input your values for row {i} - ')
                        r = t.split(',')
                        if len(valuesplit1) == len(r):
                            for j in r:
                                if j.isalpha() or j.isdigit():
                                    pass
                                else:
                                    raise ValueError
                        else:
                            raise TypeError
                        return r
                    except TypeError:
                        print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
                    except ValueError:
                        print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')

            r = tried()
            records1.append(r)
        return records1
    records1 = error_management4()
    print()

    # adding records to the table
    def adding_records_in_the_table1():
        table = doc.add_table(rows=1, cols=cols1)
        hdr_Cells = table.rows[0].cells
        for j in range(cols1):
            hdr_Cells[j].text = valuesplit1[j]
        table.style = 'Table Grid'
        # putting the values in the table
        for e in range(rows1):
            row_Cells = table.add_row().cells
            for b in range(len(records1[0])):
                row_Cells[b].text = records1[e][b]
        return table
    table1 = adding_records_in_the_table1()

    # adding styles to the table
    def table1_style():
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        table1.autofit = True
        for row in table1.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].style = doc.styles['Normal']
        for i in range(3):
            doc.add_paragraph(' ')
    table1_style()

    # creates a list 'element' that facilitates the formation of the dataframe required for the graph
    def Element1():
        element1 = []
        for v in range(len(valuesplit1)):
            o_v1 = []
            element1.append(o_v1)
        for el in range(len(element1)):
            for sublist in records1:
                element1[el].append(sublist[el])
        return element1
    element1 = Element1()

    # creating the dataframe
    def dataframe1():
        data1 = {}
        for h in range(cols1):
            key1 = valuesplit1[h]
            value2 = element1[h]
            data1[key1] = value2
        print(data1)
        return data1
    data1 = dataframe1()
    df1 = pd.DataFrame(data1)
    print(Fore.YELLOW + Style.BRIGHT + 'Graph - 1')

    # taking input for axis components
    def error_management5():
        while True:
            try:
                x_y1 = input(Fore.MAGENTA + Style.BRIGHT + '''define x and y axis only write the position of column
variables when they were defined - ''')
                valuesplit2 = x_y1.split(',')
                if len(valuesplit2) == 2:
                    for i in valuesplit2:
                        if i.isdigit():
                            pass
                        else:
                            raise ValueError
                    return valuesplit2
                else:
                    raise TypeError
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    valuesplit2 = error_management5()
    x_1 = valuesplit1[int(valuesplit2[0])]
    y_1 = valuesplit1[int(valuesplit2[1])]

    # creating the graph
    def create_graph1():
        TofGraph1 = x_1 + ' Vs '
        for tofg in range(1, len(valuesplit2)):
            TofGraph1 = TofGraph1 + valuesplit1[int(valuesplit2[tofg])]
        years = data1[x_1]
        sales = [int(s) for s in data1[y_1]]
        plt.figure(figsize=(8, 6))
        bars = plt.bar(years, sales, color='black')
        plt.xlabel(x_1)
        plt.ylabel(y_1)
        plt.title(TofGraph1)

        # Adding value labels above the bars
        for bar, sale in zip(bars, sales):
            plt.text(bar.get_x() + bar.get_width() / 2 - 0.1, bar.get_height() + 20, str(sale), fontsize=12,
                     ha='center')

        plt.savefig(f'', dpi=600, bbox_inches='tight',  #add file path in this format (keep the variable) C:\\Users\\OneDrive\\Desktop\\code\\{TofGraph1}.png
                    pad_inches=0.2)  # file path

        doc.add_picture(f'', width=Cm(16.51),  #add file path in this format (keep the variable) C:\\Users\\OneDrive\\Desktop\\code\\{TofGraph1}.png
                        height=Inches(4))
        doc.add_page_break()
    create_graph1()

    table_str1 = df1.to_string(index=False) # for gpt to analyse the graph we are converting the dataframe to a string

                                            # Para 3 - ANALYSIS
    # getting response from ChatGPT
    def chatgpt2():
        # Convert the DataFrame to a string
        prompt = f"create an elaborate analysis for the employees performance using {table_str1} for the company named {H1}, considering the values mentioned are in rupees. Dont add any title to the response"
        completion = openai.Completion.create(engine='text-davinci-003', prompt=prompt, max_tokens=1024, n=3, stop=None, temperature=1)  # controls the level of randomness in the output. The higher the temperature, the more varied and less coherent response output
        response2 = completion.choices[0].text
        return response2
    response2 = chatgpt2()

    # adding styles to para 3
    def p3_style():
        p3 = doc.add_paragraph(response2)
        font3 = p3.style.font
        font3.name = 'Cascadia Code'
        font3.size = Pt(12)
        font3.color.rgb = RGBColor(255, 255, 255)
        p3_format = p3.paragraph_format
        p3_format.space_before = Pt(0)
        p3_format.space_after = Pt(0)
        p3_format.keep_together = True
        p3_format.keep_with_next = False
        p3_format.page_break_before = False
        p3_format.widow_control = False
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_page_break()
    p3_style()
    return table_str1
table_str1 = create_H2_Para2()
# Heading2 function complete


                                        # Heading 3 - SALES TEAM PERFORMANCE

# creating section 3
def create_H3_Para3():

    # adding styles to H3
    def H3_style():
        H3 = 'Sales team Performance'
        head3 = doc.add_heading(H3, 6)
        font3 = head3.style.font
        font3.size = Pt(48)
        font3.bold = True
        font3.italic = False
        font3.color.rgb = RGBColor(160, 160, 160)
        head3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(2):
            doc.add_paragraph(' ')
    H3_style()
    print()
    print(Fore.YELLOW + Style.BRIGHT + 'Table - 2')

                                             # Para 3[Table and Graph]
    # taking input about the number of rows and columns
    def error_management6():
        while True:
            try:
                row_col2 = input(Fore.MAGENTA + Style.BRIGHT + 'How many rows and columns do you want - ')  # taking the order of the table
                new_row_col2 = row_col2.split(',')
                if len(new_row_col2) == 2:
                    for i in new_row_col2:
                        if i.isdigit():
                            pass
                        else:
                            raise ValueError
                else:
                    raise TypeError
                return new_row_col2
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    new_row_col2 = error_management6()
    rows2 = int(new_row_col2[0])
    cols2 = int(new_row_col2[1])
    print()

    # taking input of the column titles
    def error_management7():
        while True:
            try:
                value3 = input(Fore.MAGENTA + Style.BRIGHT + 'Input the column Titles - ')
                valuesplit3 = value3.split(',')
                if len(valuesplit3) == cols2:
                    for i in valuesplit3:
                        if i.replace(' ', '').replace('.', '').replace('%', '').replace('(', '').replace(')', '').replace('$', '').isalpha():
                            pass
                        else:
                            raise ValueError
                else:
                    raise TypeError
                return valuesplit3
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    valuesplit3 = error_management7()
    print()

    # taking table values from the user
    def error_management8():
        records2 = []
        for i in range(1, rows2 + 1):
            def tried():
                while True:
                    try:
                        t = input(Fore.MAGENTA + Style.BRIGHT + f'Input your values for row {i} - ')
                        r = t.split(',')
                        if len(valuesplit3) == len(r):
                            for j in r:
                                if j.isalpha() or j.isdigit():
                                    pass
                                else:
                                    raise ValueError
                        else:
                            raise TypeError
                        return r
                    except TypeError:
                        print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
                    except ValueError:
                        print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')

            r = tried()
            records2.append(r)
        return records2
    records2 = error_management8()
    print()

    # adding records to the table
    def adding_records_in_the_table2():
        table = doc.add_table(rows=1, cols=cols2)
        hdr_Cells = table.rows[0].cells
        for j in range(cols2):
            hdr_Cells[j].text = valuesplit3[j]
        table.style = 'Table Grid'
        # putting the values in the table
        for e in range(rows2):
            row_Cells = table.add_row().cells
            for b in range(len(records2[0])):
                row_Cells[b].text = records2[e][b]
        return table
    table2 = adding_records_in_the_table2()

    # adding styles to the table
    def table2_style():
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.autofit = True
        for row in table2.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].style = doc.styles['Normal']
    table2_style()

    # creates a list which facilitates the formation of the dataframe required for creating the bar graph
    def Element2():
        for i in range(3):
            doc.add_paragraph(' ')
        element2 = []
        for v in range(len(valuesplit3)):
            o_v2 = []
            element2.append(o_v2)
        for el in range(len(element2)):
            for sublist in records2:
                element2[el].append(sublist[el])
        return element2
    element2 = Element2()

    # creating the dataframe
    def dataframe2():
        data2 = {}
        for h in range(cols2):
            key2 = valuesplit3[h]
            value4 = element2[h]
            data2[key2] = value4
        return data2
    data2 = dataframe2()
    df2 = pd.DataFrame(data2)
    print(Fore.YELLOW + Style.BRIGHT + 'Graph - 2')

    # taking axis components from the user
    def error_management9():
        while True:
            try:
                x_y2 = input(Fore.MAGENTA + Style.BRIGHT + '''define x and y axis only write the position of column
variables when they were defined - ''')
                valuesplit4 = x_y2.split(',')
                if len(valuesplit4) == 2:
                    for i in valuesplit4:
                        if i.isdigit():
                            pass
                        else:
                            raise ValueError
                    return valuesplit4
                else:
                    raise TypeError
            except ValueError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
            except TypeError:
                print(Fore.RED + Style.BRIGHT + 'ERROR, please enter again')
    valuesplit4 = error_management9()
    x_2 = valuesplit3[int(valuesplit4[0])]
    y_2 = valuesplit3[int(valuesplit4[1])]
    plt.clf()

    # creating the graph
    def create_graph2():
        TofGraph2 = x_2 + ' Vs '
        for tofg in range(1, len(valuesplit4)):
            TofGraph2 = TofGraph2 + valuesplit3[int(valuesplit4[tofg])]
        employee = data2[x_2]
        sales = [int(s) for s in data2[y_2]]
        plt.figure(figsize=(8, 6))
        bars = plt.bar(employee, sales, color='black')
        plt.xlabel(x_2)
        plt.ylabel(y_2)
        plt.title(TofGraph2)

        # Adding value labels above the bars
        for bar, sale in zip(bars, sales):
            max_label_height = 10  # Setting the  maximum label height
            y_coordinate = min(bar.get_height() + 20, max_label_height)
            plt.text(bar.get_x() + bar.get_width() / 2 - 0.1, y_coordinate, str(sale), fontsize=12, ha='center')

        plt.savefig(f'', dpi=600, bbox_inches='tight', #add file path in this format (keep the variable) C:\\Users\\OneDrive\\Desktop\\code\\{TofGraph2}.png
                    pad_inches=0.2)  # file path

        doc.add_picture(f'', width=Cm(16.51), #add file path in this format (keep the variable) C:\\Users\\OneDrive\\Desktop\\code\\{TofGraph1}.png
                        height=Inches(4))
        doc.add_page_break()

    create_graph2()

    table_str2 = df2.to_string(index=False)  # for gpt to analyse the graph we are converting the dataframe to a string

    # getting response from ChatGPT
    def chatgpt3():
        # Convert the DataFrame to a string
        prompt = f" create an elaborate employee performance analysis report in a paragraph form using the table {table_str2} for the company named {H1}, considering the values mentioned are in rupees. Also dont add any title to the response."
        completion = openai.Completion.create(engine='text-davinci-003', prompt=prompt, max_tokens=1024,  n=3, stop=None,temperature=0)  # controls the level of randomness in the output. The higher the temperature, the more varied and less coherent response output
        response3 = completion.choices[0].text
        return response3
    response3 = chatgpt3()

    # adding style to para 4
    def p4_style():
        p4 = doc.add_paragraph(response3)
        font4 = p4.style.font
        font4.name = 'Cascadia Code'
        font4.size = Pt(12)
        font4.color.rgb = RGBColor(255, 255, 255)
        p4_format = p4.paragraph_format
        p4_format.space_before = Pt(0)
        p4_format.space_after = Pt(0)
        p4_format.keep_together = True
        p4_format.keep_with_next = False
        p4_format.page_break_before = False
        p4_format.widow_control = False
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_page_break()
    p4_style()
    return table_str2
table_str2 = create_H3_Para3()
# Heading3 function complete

                                            # Heading 4 - IMPROVEMENTS

# creating section 4
def create_H4_Para4():

    # adding styles to H4
    def H4_style():
        H4 = 'Improvements '
        head4 = doc.add_heading(H4, 6)
        font4 = head4.style.font
        font4.size = Pt(48)
        font4.bold = True
        font4.italic = False
        font4.color.rgb = RGBColor(160, 160, 160)
        head4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    H4_style()

                                             # Para 5 - SUGGESTIONS
    # getting from ChatGPT
    def chatgpt4():
        prompt = f"Suggest improvements that can be made according the the annual sales report {table_str2} for the company named {H1} in bullet points."
        completion = openai.Completion.create(engine='text-davinci-003', prompt=prompt, max_tokens=1024,  n=3, stop=None,temperature=0)  # controls the level of randomness in the output. The higher the temperature, the more varied and less coherent response output
        response4 = completion.choices[0].text
        return response4
    response4 = chatgpt4()

    # adding style to para 5
    def p5_style():
        p5 = doc.add_paragraph(response4)
        font5 = p5.style.font
        font5.name = 'Cascadia Code'
        font5.size = Pt(12)
        font5.color.rgb = RGBColor(255, 255, 255)
        p5_format = p5.paragraph_format
        p5_format.space_before = Pt(0)
        p5_format.space_after = Pt(0)
        p5_format.keep_together = True
        p5_format.keep_with_next = False
        p5_format.page_break_before = False
        p5_format.widow_control = False
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p5_style()
create_H4_Para4()
# Heading4 function complete

add_background()
create_sections()

file_name = f''  # file path in this format (keep this variable) C:\\Users\\OneDrive\\Desktop\\code\\{H1} - Annual Performance.docx
doc.save(file_name) # saving file
os.startfile(file_name) # opening file automatically
